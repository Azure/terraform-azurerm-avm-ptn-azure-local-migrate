"""Generate the Azure Migrate for Terraform user guide as a .docx file."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "Azure-Migrate-Terraform-User-Guide.docx"

CODE_FONT = "Consolas"
CODE_SHADE = "F2F2F2"


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    shade_cell(cell, CODE_SHADE)
    cell.paragraphs[0].text = ""
    first = True
    for line in code.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else "")
        run.font.name = CODE_FONT
        run.font.size = Pt(9)
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
    doc.add_paragraph()


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        # support simple "code: rest" inline emphasis on leading token in backticks
        parts = item.split("`")
        for idx, part in enumerate(parts):
            run = p.add_run(part)
            if idx % 2 == 1:
                run.font.name = CODE_FONT
                run.font.size = Pt(10)


def add_note(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(text)


def build():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Title
    title = doc.add_heading("Azure Migrate for Terraform – User Guide", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    r = p.add_run(
        "\u26a0 Preview Version – This module is currently in preview. "
        "Features and APIs may change before the stable release."
    )
    r.bold = True

    # Overview
    doc.add_heading("Overview", level=1)
    add_para(
        doc,
        "The Azure Migrate for Terraform module provides infrastructure as code (IaC) "
        "capabilities for Azure Migrate, enabling automated discovery, replication, and "
        "migration of virtual machines to Azure Local. This module supports:",
    )
    add_bullets(doc, ["VMware to Azure Local migrations", "Hyper-V to Azure Local migrations"])

    doc.add_heading("Key Features", level=2)
    add_bullets(
        doc,
        [
            "Discovery: Automatically discover VMs from VMware or Hyper-V environments",
            "Infrastructure Setup: Initialize replication vaults, policies, and extensions",
            "Replication: Enable VM replication with granular disk and NIC configuration",
            "Monitoring: Track replication jobs and protected item status",
            "Migration: Perform planned failovers (migrations) to Azure Local",
            "Cleanup: Remove replication when no longer needed",
        ],
    )

    # Prerequisites
    doc.add_heading("Prerequisites", level=1)
    doc.add_heading("Required Software", level=2)
    add_table(doc, ["Component", "Version"], [["Terraform", ">= 1.9"], ["AzAPI Provider", "~> 2.4"]])

    doc.add_heading("Azure Requirements", level=2)
    add_bullets(
        doc,
        [
            "Azure Subscription with appropriate permissions",
            "Resource Group where resources will be deployed",
            "Azure Migrate Project (can be created by this module)",
            "Azure Local Cluster with: custom location configured, logical networks set up, storage paths defined",
            "Source Environment: VMware vCenter or Hyper-V host with an Azure Migrate appliance deployed and registered",
        ],
    )

    doc.add_heading("Required Permissions", level=2)
    add_para(doc, "The executing identity needs the following permissions:")
    add_bullets(
        doc,
        [
            "Contributor on the resource group",
            "User Access Administrator (for role assignments)",
            "Azure Migrate Assessor (automatically assigned to project managed identity)",
        ],
    )

    # Quick Start
    doc.add_heading("Quick Start", level=1)
    doc.add_heading("1. Configure the Provider", level=2)
    add_code_block(
        doc,
        'terraform {\n'
        '  required_version = ">= 1.9"\n\n'
        '  required_providers {\n'
        '    azapi = {\n'
        '      source  = "Azure/azapi"\n'
        '      version = "~> 2.4"\n'
        '    }\n'
        '  }\n'
        '}\n\n'
        'provider "azapi" {}',
    )

    doc.add_heading("2. Add the Module from the Registry", level=2)
    add_code_block(
        doc,
        'module "migrate" {\n'
        '  source  = "Azure/avm-ptn-azure-local-migrate/azurerm"\n'
        '  version = "0.1.2" # Pin to the latest compatible version\n\n'
        '  # ... configuration ...\n'
        '}',
    )
    add_para(
        doc,
        "Find the latest version at: "
        "https://registry.terraform.io/modules/Azure/avm-ptn-azure-local-migrate/azurerm/latest",
    )

    doc.add_heading("3. Authenticate", level=2)
    add_code_block(doc, 'az login\naz account set --subscription "<your-subscription-id>"')

    doc.add_heading("4. Use the Module", level=2)
    add_para(
        doc,
        "Each step in the migration requires a separate module invocation with a different "
        "operation_mode. See the detailed sections below.",
    )

    # Operation Modes
    doc.add_heading("Operation Modes", level=1)
    add_para(doc, "The module's behavior is controlled by the operation_mode variable. Valid values are:")
    add_table(
        doc,
        ["Mode", "Description"],
        [
            ["create-project", "Create a new Azure Migrate project with all solutions"],
            ["discover", "List discovered VMs from the source environment"],
            ["initialize", "Set up replication vault, policy, extension, and storage"],
            ["replicate", "Start replication for a specific VM"],
            ["jobs", "Query replication job status"],
            ["list", "List all protected (replicating) items in a vault"],
            ["get", "Get detailed status of a specific protected item"],
            ["migrate", "Perform planned failover (final migration)"],
            ["remove", "Disable and remove replication for a protected item"],
        ],
    )
    add_note(
        doc,
        "Note",
        "The source environment is selected with source_machine_type (\"VMware\" or \"HyperV\"). "
        "The module derives the replication instanceType (VMwareToAzStackHCI / HyperVToAzStackHCI) "
        "internally — there is no instance_type input.",
    )

    # 1. Create Project
    doc.add_heading("1. Create Project", level=2)
    add_para(
        doc,
        "Creates a new Azure Migrate project with the required solutions (Assessment, Discovery, "
        "Migration, DataReplication) and assigns the necessary role to the project's managed identity.",
    )
    add_code_block(
        doc,
        'module "create_project" {\n'
        '  source  = "Azure/avm-ptn-azure-local-migrate/azurerm"\n'
        '  version = "0.1.2"\n\n'
        '  location               = "eastus"\n'
        '  name                   = "create-project"\n'
        '  parent_id              = "/subscriptions/<sub-id>/resourceGroups/<rg-name>"\n'
        '  create_migrate_project = true\n'
        '  operation_mode         = "create-project"\n'
        '  project_name           = "my-migrate-project"\n'
        '  connectivity_method    = "Public-endpoint" # or "Private-endpoint"\n'
        '  tags = {\n'
        '    Environment = "Production"\n'
        '  }\n'
        '}',
    )
    add_para(doc, "Key inputs:", bold=True)
    add_bullets(
        doc,
        [
            "create_migrate_project — Must be true",
            "project_name — The name for the new Azure Migrate project",
            "location — Required in this mode (there is no existing project to read it from)",
            "connectivity_method — \"Public-endpoint\" (default) or \"Private-endpoint\"",
        ],
    )
    add_para(doc, "Key outputs:", bold=True)
    add_bullets(doc, ["migrate_project_id — The ARM resource ID of the created project"])
    add_note(
        doc,
        "Note",
        "Skip this step if you already have an Azure Migrate project. The module works with existing "
        "projects, and in all other modes location is read from the project when omitted.",
    )

    # 2. Discover
    doc.add_heading("2. Discover Servers", level=2)
    add_para(
        doc,
        "Queries Azure Migrate to list all discovered VMs in your source environment. Read-only; "
        "retrieves machine name, IP addresses, OS, boot type, and disk IDs.",
    )
    add_code_block(
        doc,
        'module "discover" {\n'
        '  source  = "Azure/avm-ptn-azure-local-migrate/azurerm"\n'
        '  version = "0.1.2"\n\n'
        '  name                = "discover"\n'
        '  parent_id           = "/subscriptions/<sub-id>/resourceGroups/<rg-name>"\n'
        '  operation_mode      = "discover"\n'
        '  project_name        = "my-migrate-project"\n'
        '  source_machine_type = "VMware" # or "HyperV"\n'
        '}',
    )
    add_para(doc, "Optional inputs:", bold=True)
    add_bullets(
        doc,
        [
            "appliance_name — Filter by a specific appliance (site). If omitted, all VMs in the project are returned.",
            "source_machine_type — \"VMware\" (default) or \"HyperV\"",
        ],
    )
    add_para(doc, "Key outputs:", bold=True)
    add_bullets(
        doc,
        [
            "discovered_servers — Filtered list (index, machine name, IPs, OS, boot type, OS disk ID)",
            "discovered_servers_count — Number of discovered servers with discovery data",
            "total_machines_count — Total machines including those without discovery data",
            "discovered_servers_raw — Full API response for debugging",
        ],
    )
    add_para(doc, "Example output:", bold=True)
    add_code_block(
        doc,
        'discovered_servers = [\n'
        '  {\n'
        '    index            = 1\n'
        '    machine_name     = "web-server-01"\n'
        '    ip_addresses     = ["10.0.1.10", "10.0.1.11"]\n'
        '    operating_system = "Windows Server 2019"\n'
        '    boot_type        = "EFI"\n'
        '    os_disk_id       = "6000C290-a4d0-e5ea-bad5-4e993df22e3b"\n'
        '  },\n'
        '  # ...\n'
        ']',
    )
    add_note(doc, "Tip", "Use the os_disk_id from discovery output as input for the replicate step.")

    # 3. Initialize
    doc.add_heading("3. Initialize Replication Infrastructure", level=2)
    add_para(
        doc,
        "Sets up all infrastructure required before replication can begin. Safe to re-run: if "
        "replication was already configured it validates that the project is in a good state (vault, "
        "extension, fabrics already exist) without overwriting.",
    )
    add_code_block(
        doc,
        'module "initialize" {\n'
        '  source  = "Azure/avm-ptn-azure-local-migrate/azurerm"\n'
        '  version = "0.1.2"\n\n'
        '  name                = "initialize"\n'
        '  parent_id           = "/subscriptions/<sub-id>/resourceGroups/<rg-name>"\n'
        '  operation_mode      = "initialize"\n'
        '  project_name        = "my-migrate-project"\n'
        '  source_machine_type = "VMware"\n\n'
        '  # Appliance names — fabrics are auto-discovered from these\n'
        '  source_appliance_name = "src-appliance"\n'
        '  target_appliance_name = "tgt-appliance"\n\n'
        '  # Optional: replication policy overrides (defaults shown)\n'
        '  replication_policy = {\n'
        '    recovery_point_history_minutes     = 4320 # 72 hours\n'
        '    crash_consistent_frequency_minutes = 60   # 1 hour\n'
        '    app_consistent_frequency_minutes   = 240  # 4 hours\n'
        '    # name = "custom-policy-name"             # auto-generated if omitted\n'
        '  }\n\n'
        '  # Optional: use an existing cache storage account\n'
        '  # cache_storage_account_id = "/subscriptions/.../storageAccounts/existing-account"\n'
        '}',
    )
    add_note(
        doc,
        "Fabric auto-discovery",
        "Provide source_appliance_name and target_appliance_name and the module discovers the "
        "corresponding replication fabrics automatically. There are no source_fabric_id / "
        "target_fabric_id inputs; the resolved IDs are exposed as outputs.",
    )
    add_para(doc, "Key outputs:", bold=True)
    add_bullets(
        doc,
        [
            "replication_vault_id — Vault ARM ID",
            "replication_policy_id — Policy ARM ID",
            "replication_extension_id / replication_extension_name — Extension ARM ID / name",
            "cache_storage_account_id / cache_storage_account_name — Cache storage account",
            "source_fabric_id / target_fabric_id — Resolved fabric IDs",
            "source_fabric_discovered / target_fabric_discovered — Discovered fabric details",
            "replication_fabrics_available — All fabrics in the resource group (troubleshooting)",
            "replication_vault_identity — Managed identity principal ID of the vault",
        ],
    )
    add_note(
        doc,
        "Brownfield",
        "If replication infrastructure already exists, the module detects it and skips creation — "
        "required settings are validated instead of recreated.",
    )

    # 4. Replicate
    doc.add_heading("4. Replicate VMs", level=2)
    add_para(
        doc,
        "Starts replication of a source VM to Azure Local by creating a \"protected item\" in the "
        "replication vault.",
    )
    doc.add_heading("Default User Mode (single OS disk / single NIC)", level=3)
    add_code_block(
        doc,
        'module "replicate" {\n'
        '  source  = "Azure/avm-ptn-azure-local-migrate/azurerm"\n'
        '  version = "0.1.2"\n\n'
        '  name                = "replicate-vm"\n'
        '  parent_id           = "/subscriptions/<sub-id>/resourceGroups/<rg-name>"\n'
        '  operation_mode      = "replicate"\n'
        '  project_name        = "my-migrate-project"\n'
        '  source_machine_type = "VMware"\n\n'
        '  # VM to replicate\n'
        '  machine_id = "/subscriptions/.../machines/<machine-guid>"\n\n'
        '  # OS disk (from discover output) — size is derived automatically\n'
        '  os_disk_id = "6000C290-a4d0-e5ea-bad5-4e993df22e3b"\n\n'
        '  # Network (simple mode — single target logical network)\n'
        '  target_virtual_switch_id = "/subscriptions/.../logicalnetworks/my-lnet"\n'
        '  # target_test_virtual_switch_id = "/subscriptions/.../logicalnetworks/test-lnet" # optional\n\n'
        '  # Target VM configuration\n'
        '  target_vm_name           = "my-vm-migrated"\n'
        '  target_resource_group_id = "/subscriptions/.../resourceGroups/target-rg"\n'
        '  target_storage_path_id   = "/subscriptions/.../storagecontainers/my-storage"\n'
        '  target_hci_cluster_id    = "/subscriptions/.../clusters/my-cluster"\n'
        '  custom_location_id       = "/subscriptions/.../customLocations/my-cl"\n\n'
        '  # Optional: target VM compute overrides (defaults shown)\n'
        '  target_vm_compute = {\n'
        '    cpu_cores                 = 2\n'
        '    ram_mb                    = 4096\n'
        '    is_dynamic_memory_enabled = false\n'
        '    hyperv_generation         = "1" # "1" or "2"\n'
        '  }\n\n'
        '  # Appliance info — replication vault, extension, policy, fabrics,\n'
        '  # and run-as account are auto-resolved from the project/appliances.\n'
        '  source_appliance_name = "src-appliance"\n'
        '  target_appliance_name = "tgt-appliance"\n\n'
        '  # Optional explicit overrides:\n'
        '  # replication_vault_id = module.initialize.replication_vault_id\n'
        '  # run_as_account_id    = "/subscriptions/.../runasaccounts/<account-id>"\n'
        '}',
    )
    doc.add_heading("Power User Mode (multiple disks / NICs)", level=3)
    add_para(doc, "For VMs with multiple disks or NICs, use disks_to_include and nics_to_include:")
    add_code_block(
        doc,
        'module "replicate" {\n'
        '  source  = "Azure/avm-ptn-azure-local-migrate/azurerm"\n'
        '  version = "0.1.2"\n\n'
        '  # ... (common configuration as above, minus os_disk_id / target_virtual_switch_id) ...\n\n'
        '  disks_to_include = [\n'
        '    {\n'
        '      disk_id          = "6000C290-a4d0-e5ea-bad5-4e993df22e3b"\n'
        '      disk_size_gb     = 40\n'
        '      disk_file_format = "VHDX"\n'
        '      is_os_disk       = true\n'
        '      is_dynamic       = true\n'
        '    },\n'
        '    {\n'
        '      disk_id          = "6000C29a-bcb7-a62b-7ed0-0c78f3dc1f80"\n'
        '      disk_size_gb     = 100\n'
        '      disk_file_format = "VHDX"\n'
        '      is_os_disk       = false\n'
        '      is_dynamic       = true\n'
        '    }\n'
        '  ]\n\n'
        '  nics_to_include = [\n'
        '    {\n'
        '      nic_id            = "4000"\n'
        '      target_network_id = "/subscriptions/.../logicalnetworks/prod-lnet"\n'
        '      test_network_id   = "/subscriptions/.../logicalnetworks/test-lnet"\n'
        '      selection_type    = "SelectedByUser"\n'
        '    }\n'
        '  ]\n'
        '}',
    )
    add_para(doc, "Key outputs:", bold=True)
    add_bullets(
        doc,
        [
            "protected_item_id — ARM ID of the created protected item (needed for get/migrate/remove)",
            "protected_item_name — Name of the protected item",
            "replication_state — Current replication health",
        ],
    )

    # 5. Jobs
    doc.add_heading("5. Monitor Replication Jobs", level=2)
    add_para(doc, "Queries replication job status. List all jobs or retrieve a specific job by name.")
    add_code_block(
        doc,
        '# List all jobs\n'
        'module "jobs" {\n'
        '  source  = "Azure/avm-ptn-azure-local-migrate/azurerm"\n'
        '  version = "0.1.2"\n\n'
        '  name           = "replication-jobs"\n'
        '  parent_id      = "/subscriptions/<sub-id>/resourceGroups/<rg-name>"\n'
        '  operation_mode = "jobs"\n'
        '  project_name   = "my-migrate-project"\n'
        '  # Optional — auto-resolved from the project when omitted:\n'
        '  # replication_vault_id = "/subscriptions/.../replicationVaults/my-vault"\n'
        '}\n\n'
        '# Get a specific job\n'
        'module "specific_job" {\n'
        '  source  = "Azure/avm-ptn-azure-local-migrate/azurerm"\n'
        '  version = "0.1.2"\n\n'
        '  # ... same as above, plus:\n'
        '  job_name = "my-specific-job-name"\n'
        '}',
    )
    add_para(doc, "Key outputs:", bold=True)
    add_bullets(
        doc,
        [
            "replication_jobs — Summary of all jobs (name, state, VM, timestamps)",
            "replication_jobs_count — Total number of jobs",
            "replication_job — Detailed info for a specific job (tasks and errors)",
            "vault_id_for_jobs — Vault ID used for the query",
        ],
    )

    # 6. List
    doc.add_heading("6. List Protected Items", level=2)
    add_para(doc, "Lists all protected (replicating) items in a replication vault.")
    add_code_block(
        doc,
        'module "list" {\n'
        '  source  = "Azure/avm-ptn-azure-local-migrate/azurerm"\n'
        '  version = "0.1.2"\n\n'
        '  name           = "list-items"\n'
        '  parent_id      = "/subscriptions/<sub-id>/resourceGroups/<rg-name>"\n'
        '  operation_mode = "list"\n'
        '  project_name   = "my-migrate-project" # vault auto-discovered\n'
        '}',
    )
    add_para(doc, "Key outputs:", bold=True)
    add_bullets(
        doc,
        [
            "protected_items_summary — Summary (name, ID, state, health, source/target VM names)",
            "protected_items_count — Total number of protected items",
            "protected_items_list — Complete raw list",
            "protected_items_by_state — Items grouped by protection state",
            "protected_items_by_health — Items grouped by replication health",
            "protected_items_with_errors — Items that have health errors",
        ],
    )

    # 7. Get
    doc.add_heading("7. Get Protected Item Details", level=2)
    add_para(doc, "Retrieves detailed status for a specific protected item, by ID or by name.")
    add_code_block(
        doc,
        '# By ID\n'
        'module "get_by_id" {\n'
        '  source  = "Azure/avm-ptn-azure-local-migrate/azurerm"\n'
        '  version = "0.1.2"\n\n'
        '  name              = "get-item"\n'
        '  parent_id         = "/subscriptions/<sub-id>/resourceGroups/<rg-name>"\n'
        '  operation_mode    = "get"\n'
        '  project_name      = "my-migrate-project"\n'
        '  protected_item_id = "/subscriptions/.../protectedItems/<item-name>"\n'
        '}\n\n'
        '# By name\n'
        'module "get_by_name" {\n'
        '  source  = "Azure/avm-ptn-azure-local-migrate/azurerm"\n'
        '  version = "0.1.2"\n\n'
        '  # ... same as above, but use protected_item_name instead:\n'
        '  protected_item_name = "my-protected-item"\n'
        '}',
    )
    add_para(doc, "Key outputs:", bold=True)
    add_bullets(
        doc,
        [
            "protected_item — Complete protected item details",
            "protected_item_summary — Key fields: state, health, allowed jobs, source/target names",
            "protected_item_custom_properties — Fabric-specific details, disk config, network settings",
            "protected_item_health_errors — Any health errors",
        ],
    )

    # 8. Migrate
    doc.add_heading("8. Migrate (Planned Failover)", level=2)
    add_para(
        doc,
        "Performs the final migration by executing a planned failover, creating the VM on the target "
        "Azure Local cluster.",
    )
    add_code_block(
        doc,
        'module "migrate" {\n'
        '  source  = "Azure/avm-ptn-azure-local-migrate/azurerm"\n'
        '  version = "0.1.2"\n\n'
        '  name               = "migrate-vm"\n'
        '  parent_id          = "/subscriptions/<sub-id>/resourceGroups/<rg-name>"\n'
        '  operation_mode     = "migrate"\n'
        '  protected_item_id  = "/subscriptions/.../protectedItems/<item-name>"\n'
        '  shutdown_source_vm = true # Recommended for production\n'
        '}',
    )
    add_para(doc, "Key inputs:", bold=True)
    add_bullets(
        doc,
        [
            "protected_item_id — The protected item ARM ID (from replicate or list outputs)",
            "shutdown_source_vm — true to shut down the source VM before failover (recommended for data consistency)",
        ],
    )
    add_para(doc, "Key outputs:", bold=True)
    add_bullets(
        doc,
        [
            "migration_status — Operation status, source/target VM names",
            "migration_protected_item_details — Details of the item being migrated",
            "migration_validation_warnings — Health warnings detected before migration",
            "migration_operation_details — Raw async operation response (for job tracking)",
        ],
    )
    add_note(
        doc,
        "Important",
        "The protected item must be in a Protected state with PlannedFailover in its allowedJobs "
        "before migration can proceed. Use the get operation to verify readiness.",
    )

    # 9. Remove
    doc.add_heading("9. Remove Replication", level=2)
    add_para(doc, "Disables and removes replication for a protected item.")
    add_code_block(
        doc,
        'module "remove" {\n'
        '  source  = "Azure/avm-ptn-azure-local-migrate/azurerm"\n'
        '  version = "0.1.2"\n\n'
        '  name             = "remove-replication"\n'
        '  parent_id        = "/subscriptions/<sub-id>/resourceGroups/<rg-name>"\n'
        '  operation_mode   = "remove"\n'
        '  target_object_id = "/subscriptions/.../protectedItems/<item-name>"\n'
        '  force_remove     = false # Set true only if normal removal fails\n'
        '}',
    )
    add_para(doc, "Key inputs:", bold=True)
    add_bullets(
        doc,
        [
            "target_object_id — The protected item ARM ID to remove",
            "force_remove — Force removal even if the item is in an inconsistent state",
        ],
    )
    add_para(doc, "Key outputs:", bold=True)
    add_bullets(
        doc,
        [
            "removal_status — Operation status and confirmation message",
            "protected_item_details — Details of the item before removal",
            "removal_operation_headers — Async operation headers for job tracking",
        ],
    )
    add_note(
        doc,
        "Caution",
        "Setting force_remove = true may leave resources in an inconsistent state. Use it only as a "
        "last resort.",
    )

    # 10. End-to-End
    doc.add_heading("10. End-to-End (Replicate + Migrate Together)", level=2)
    add_para(
        doc,
        "The examples/end-to-end/ example demonstrates a complete workflow for multiple VMs in a "
        "single configuration:",
    )
    add_code_block(
        doc,
        "Step 0: Initialize -> Step 1: Replicate -> Wait for sync -> Step 2: Verify -> Step 3: Migrate",
    )
    add_para(doc, "Usage", bold=True)
    add_bullets(
        doc,
        [
            "Copy the examples/end-to-end/ directory.",
            "Create terraform.tfvars from the example: cp terraform.tfvars.example terraform.tfvars",
            "Edit terraform.tfvars with your environment values.",
            "Define your VMs in the vms variable map (see below).",
            "Run: terraform init then terraform apply.",
        ],
    )
    add_code_block(
        doc,
        'vms = {\n'
        '  "web-server" = {\n'
        '    machine_id       = "/subscriptions/.../machines/<guid>"\n'
        '    target_vm_name   = "web-server-migrated"\n'
        '    os_disk_id       = "6000C290-xxxx-xxxx-xxxx-xxxxxxxxxxxx"\n'
        '    disks_to_include = [\n'
        '      {\n'
        '        disk_id      = "6000C290-xxxx-xxxx-xxxx-xxxxxxxxxxxx"\n'
        '        disk_size_gb = 40\n'
        '        is_os_disk   = true\n'
        '      }\n'
        '    ]\n'
        '    nics_to_include = [\n'
        '      {\n'
        '        nic_id            = "4000"\n'
        '        target_network_id = "/subscriptions/.../logicalnetworks/my-lnet"\n'
        '      }\n'
        '    ]\n'
        '  }\n'
        '}',
    )
    add_para(
        doc,
        "The example includes a PowerShell polling step that waits for initial replication to complete "
        "before migrating. This can take minutes to hours depending on VM disk size and network "
        "bandwidth.",
    )

    # Configuration Modes
    doc.add_heading("Configuration Modes", level=1)
    doc.add_heading("Default User Mode", level=2)
    add_para(doc, "For VMs with a single OS disk and a single target network:")
    add_table(
        doc,
        ["Variable", "Description"],
        [
            ["os_disk_id", "OS disk identifier from discovery (size auto-derived)"],
            ["target_virtual_switch_id", "Target logical network ARM ID"],
            ["target_test_virtual_switch_id", "(Optional) Test logical network ARM ID"],
        ],
    )
    doc.add_heading("Power User Mode", level=2)
    add_para(doc, "For VMs with multiple disks and/or NICs:")
    add_table(
        doc,
        ["Variable", "Description"],
        [
            ["disks_to_include", "List of disk objects (disk_id, disk_size_gb, disk_file_format, is_os_disk, is_dynamic)"],
            ["nics_to_include", "List of NIC objects (nic_id, target_network_id, test_network_id, selection_type)"],
        ],
    )
    add_para(
        doc,
        "Power user mode takes priority — when disks_to_include is non-empty, os_disk_id is ignored; "
        "likewise nics_to_include overrides target_virtual_switch_id.",
    )

    # Supported Migration Paths
    doc.add_heading("Supported Migration Paths", level=1)
    add_table(
        doc,
        ["Source", "Target", "source_machine_type", "Derived instanceType"],
        [
            ["VMware vCenter", "Azure Local", "VMware", "VMwareToAzStackHCI"],
            ["Hyper-V", "Azure Local", "HyperV", "HyperVToAzStackHCI"],
        ],
    )
    add_para(
        doc,
        "Set source_machine_type to match your source environment. This controls fabric discovery, "
        "replication extension configuration, and the failover API used.",
    )

    # Variable Reference
    doc.add_heading("Variable Reference", level=1)
    doc.add_heading("Required Variables", level=2)
    add_table(
        doc,
        ["Variable", "Type", "Description"],
        [
            ["name", "string", "Name of the migration resource (2-80 chars, alphanumeric + hyphens)"],
            ["parent_id", "string", "Resource group ARM ID (/subscriptions/{sub}/resourceGroups/{rg})"],
        ],
    )
    doc.add_heading("Operation Control", level=2)
    add_table(
        doc,
        ["Variable", "Type", "Default", "Description"],
        [
            ["operation_mode", "string", "\"discover\"", "Which operation to perform"],
            ["source_machine_type", "string", "\"VMware\"", "Source environment type (VMware / HyperV)"],
            ["create_migrate_project", "bool", "false", "Whether to create a new project"],
            ["location", "string", "null", "Azure region. Required only for create-project; otherwise read from the project"],
        ],
    )
    doc.add_heading("Project & Infrastructure", level=2)
    add_table(
        doc,
        ["Variable", "Type", "Default", "Description"],
        [
            ["project_name", "string", "null", "Azure Migrate project name"],
            ["connectivity_method", "string", "\"Public-endpoint\"", "Public-endpoint or Private-endpoint"],
            ["appliance_name", "string", "null", "Appliance name for filtering discovery"],
            ["source_appliance_name", "string", "null", "Source appliance name (for fabric auto-discovery)"],
            ["target_appliance_name", "string", "null", "Target appliance name (for fabric auto-discovery)"],
            ["cache_storage_account_id", "string", "null", "Existing cache storage account ARM ID"],
            ["replication_vault_id", "string", "null", "Vault ARM ID (auto-resolved for replicate/jobs/list/get when omitted)"],
        ],
    )
    doc.add_heading("Replication Policy (initialize)", level=2)
    add_para(doc, "replication_policy is an object with these optional fields:")
    add_table(
        doc,
        ["Field", "Type", "Default", "Description"],
        [
            ["name", "string", "auto", "Custom policy name"],
            ["recovery_point_history_minutes", "number", "4320 (72h)", "Recovery point retention"],
            ["crash_consistent_frequency_minutes", "number", "60 (1h)", "Crash-consistent snapshot interval"],
            ["app_consistent_frequency_minutes", "number", "240 (4h)", "App-consistent snapshot interval"],
        ],
    )
    doc.add_heading("VM Replication", level=2)
    add_table(
        doc,
        ["Variable", "Type", "Default", "Description"],
        [
            ["machine_id", "string", "null", "Source VM ARM ID from discovery"],
            ["machine_name", "string", "null", "Source machine internal name"],
            ["os_disk_id", "string", "null", "OS disk ID (simple mode)"],
            ["disks_to_include", "list(object)", "[]", "Disks (power user mode)"],
            ["nics_to_include", "list(object)", "[]", "NICs (power user mode)"],
            ["target_vm_name", "string", "null", "Name for the migrated VM"],
            ["run_as_account_id", "string", "null", "Run-as account ARM ID (auto-discovered when omitted)"],
        ],
    )
    add_para(doc, "target_vm_compute is an object with these optional fields:")
    add_table(
        doc,
        ["Field", "Type", "Default", "Description"],
        [
            ["cpu_cores", "number", "2", "vCPUs assigned to the migrated VM (1-240)"],
            ["ram_mb", "number", "4096", "Memory (MB) for the migrated VM"],
            ["is_dynamic_memory_enabled", "bool", "false", "Enable dynamic memory on target"],
            ["hyperv_generation", "string", "\"1\"", "Hyper-V generation (\"1\" or \"2\")"],
        ],
    )
    doc.add_heading("Target Configuration", level=2)
    add_table(
        doc,
        ["Variable", "Type", "Default", "Description"],
        [
            ["custom_location_id", "string", "null", "Azure Local custom location ARM ID"],
            ["target_hci_cluster_id", "string", "null", "Target HCI cluster ARM ID"],
            ["target_resource_group_id", "string", "null", "Target resource group ARM ID"],
            ["target_storage_path_id", "string", "null", "Target storage container ARM ID"],
            ["target_virtual_switch_id", "string", "null", "Target logical network ARM ID (simple mode)"],
            ["target_test_virtual_switch_id", "string", "null", "Test logical network ARM ID (simple mode)"],
        ],
    )
    doc.add_heading("Migration & Removal", level=2)
    add_table(
        doc,
        ["Variable", "Type", "Default", "Description"],
        [
            ["protected_item_id", "string", "null", "Protected item ARM ID (for get/migrate)"],
            ["shutdown_source_vm", "bool", "false", "Shut down source before migration"],
            ["target_object_id", "string", "null", "Protected item ARM ID (for remove)"],
            ["force_remove", "bool", "false", "Force removal of replication"],
            ["job_name", "string", "null", "Specific job name (jobs mode)"],
            ["protected_item_name", "string", "null", "Protected item name (get mode)"],
        ],
    )
    doc.add_heading("AVM Interface Variables", level=2)
    add_table(
        doc,
        ["Variable", "Type", "Default", "Description"],
        [
            ["enable_telemetry", "bool", "true", "Toggle module telemetry"],
            ["tags", "map(string)", "null", "Resource tags"],
            ["lock", "object", "null", "Resource lock (kind, name)"],
            ["managed_identities", "object", "{}", "System/user-assigned identities"],
            ["role_assignments", "map(object)", "{}", "Role assignments to create"],
            ["diagnostic_settings", "map(object)", "{}", "Diagnostic settings to create"],
        ],
    )

    # Output Reference
    doc.add_heading("Output Reference", level=1)
    doc.add_heading("Discover Mode", level=2)
    add_table(
        doc,
        ["Output", "Description"],
        [
            ["discovered_servers", "Filtered list: machine name, IPs, OS, boot type, OS disk ID"],
            ["discovered_servers_count", "Number of discovered servers with discovery data"],
            ["total_machines_count", "Total machines (including those without discovery data)"],
            ["discovered_servers_raw", "Full API response"],
        ],
    )
    doc.add_heading("Initialize Mode", level=2)
    add_table(
        doc,
        ["Output", "Description"],
        [
            ["replication_vault_id", "Vault ARM ID"],
            ["replication_vault_identity", "Vault managed identity principal ID"],
            ["replication_policy_id", "Policy ARM ID"],
            ["replication_extension_id / replication_extension_name", "Extension ARM ID / name"],
            ["cache_storage_account_id / cache_storage_account_name", "Cache storage account"],
            ["source_fabric_id / target_fabric_id", "Resolved fabric ARM IDs"],
            ["source_fabric_discovered / target_fabric_discovered", "Discovered fabric details"],
            ["replication_fabrics_available", "All fabrics in the resource group"],
        ],
    )
    doc.add_heading("Replicate Mode", level=2)
    add_table(
        doc,
        ["Output", "Description"],
        [
            ["protected_item_id", "Protected item ARM ID"],
            ["protected_item_name", "Protected item name"],
            ["replication_state", "Current replication health"],
        ],
    )
    doc.add_heading("Jobs Mode", level=2)
    add_table(
        doc,
        ["Output", "Description"],
        [
            ["replication_jobs", "Summary of all jobs"],
            ["replication_jobs_count", "Total job count"],
            ["replication_job", "Details for a specific job"],
            ["vault_id_for_jobs", "Vault ID used for the query"],
        ],
    )
    doc.add_heading("List Mode", level=2)
    add_table(
        doc,
        ["Output", "Description"],
        [
            ["protected_items_summary", "Summary of all protected items"],
            ["protected_items_list", "Complete raw list"],
            ["protected_items_count", "Total protected item count"],
            ["protected_items_by_state", "Items grouped by state"],
            ["protected_items_by_health", "Items grouped by health"],
            ["protected_items_with_errors", "Items with health errors"],
        ],
    )
    doc.add_heading("Get Mode", level=2)
    add_table(
        doc,
        ["Output", "Description"],
        [
            ["protected_item", "Complete protected item details"],
            ["protected_item_summary", "Key protected item info"],
            ["protected_item_custom_properties", "Disk/network/fabric details"],
            ["protected_item_health_errors", "Health errors"],
        ],
    )
    doc.add_heading("Migrate Mode", level=2)
    add_table(
        doc,
        ["Output", "Description"],
        [
            ["migration_status", "Operation status and VM names"],
            ["migration_protected_item_details", "Pre-migration item details"],
            ["migration_validation_warnings", "Pre-migration warnings"],
            ["migration_operation_details", "Raw async operation response"],
        ],
    )
    doc.add_heading("Remove Mode", level=2)
    add_table(
        doc,
        ["Output", "Description"],
        [
            ["removal_status", "Operation status and confirmation"],
            ["protected_item_details", "Item details before removal"],
            ["removal_operation_headers", "Async operation headers"],
        ],
    )
    doc.add_heading("Common Outputs (all modes)", level=2)
    add_table(
        doc,
        ["Output", "Description"],
        [
            ["operation_mode", "Current operation mode"],
            ["migrate_project_id / resource_id", "Azure Migrate project ARM ID"],
            ["project_name_output", "Project name"],
            ["resource_group_id", "Resource group ARM ID"],
            ["location_output", "Effective Azure region"],
            ["machine_id", "Machine ID being replicated"],
            ["target_vm_name_output", "Target VM name"],
        ],
    )

    # Brownfield
    doc.add_heading("Brownfield Support", level=1)
    add_para(
        doc,
        "The module works safely with existing Azure Migrate infrastructure, detecting and skipping "
        "creation of resources that already exist:",
    )
    add_table(
        doc,
        ["Resource", "Detection Method"],
        [
            ["Replication Vault", "Queries the DataReplication solution for an existing vaultId"],
            ["Replication Policy", "Lists existing policies and matches on instanceType"],
            ["Replication Extension", "Lists existing extensions and matches on expected naming pattern"],
            ["Role Assignments", "Lists assignments on cache storage account and checks for matching principal+role"],
            ["Cache Storage Account", "Uses cache_storage_account_id if provided, else reuses the project's recorded replicationStorageAccountId"],
        ],
    )
    add_para(
        doc,
        "You can safely run initialize multiple times — existing resources are reused, and only "
        "missing resources are created.",
    )

    # Support
    doc.add_heading("Support", level=1)
    add_bullets(
        doc,
        [
            "Terraform Registry: https://registry.terraform.io/modules/Azure/avm-ptn-azure-local-migrate/azurerm/latest",
            "AVM Documentation: https://azure.github.io/Azure-Verified-Modules/",
            "Azure Migrate Documentation: https://learn.microsoft.com/azure/migrate/",
        ],
    )

    doc.add_paragraph()
    add_para(
        doc,
        "This guide is for the preview version of Azure Migrate for Terraform. Features and APIs may "
        "change in future releases.",
        italic=True,
    )

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
